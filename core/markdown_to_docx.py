"""
Convert Markdown to DOCX format using python-docx.
"""

import re
from pathlib import Path
from typing import Optional, Union

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    raise ImportError("python-docx is required for DOCX conversion. Install with: pip install python-docx")


def convert_markdown_to_docx(markdown_content: str, output_path: Optional[str] = None) -> str:
    """
    Convert Markdown content to DOCX format.
    
    Args:
        markdown_content: The Markdown content to convert
        output_path: Path to save the DOCX output
        
    Returns:
        The DOCX content as bytes if no output_path, otherwise the output path
    """
    doc = Document()
    
    # Process Markdown content line by line
    lines = markdown_content.split('\n')
    in_code_block = False
    in_list = False
    list_items = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines unless we're in a code block
        if not line and not in_code_block:
            if in_list:
                _add_list_to_doc(doc, list_items)
                list_items = []
                in_list = False
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            if in_list:
                _add_list_to_doc(doc, list_items)
                list_items = []
                in_list = False
            heading = doc.add_heading(line[2:], level=1)
            
        elif line.startswith('## '):
            if in_list:
                _add_list_to_doc(doc, list_items)
                list_items = []
                in_list = False
            heading = doc.add_heading(line[3:], level=2)
            
        elif line.startswith('### '):
            if in_list:
                _add_list_to_doc(doc, list_items)
                list_items = []
                in_list = False
            heading = doc.add_heading(line[4:], level=3)
        
        # Images
        elif '![' in line and '](' in line:
            if in_list:
                _add_list_to_doc(doc, list_items)
                list_items = []
                in_list = False
            
            alt_text = re.search(r'!\[(.*?)\]', line)
            image_path = re.search(r'\((.*?)\)', line)
            
            if alt_text and image_path:
                alt_text = alt_text.group(1)
                image_path = image_path.group(1)
                
                # Try to add image if file exists
                if Path(image_path).exists():
                    try:
                        doc.add_picture(image_path, width=Inches(4))
                        if alt_text:
                            p = doc.add_paragraph()
                            p.add_run(f"Figure: {alt_text}").italic = True
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception:
                        # If image can't be added, add alt text
                        doc.add_paragraph(f"[Image: {alt_text}]")
                else:
                    doc.add_paragraph(f"[Image: {alt_text} - {image_path}]")
        
        # Code blocks
        elif line.startswith('```'):
            if in_list:
                _add_list_to_doc(doc, list_items)
                list_items = []
                in_list = False
                
            if in_code_block:
                in_code_block = False
            else:
                in_code_block = True
        
        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            if not in_list:
                in_list = True
            list_items.append(line.strip()[2:])
        
        # Regular text or code content
        else:
            if in_list and not line.strip().startswith('- ') and not line.strip().startswith('* '):
                _add_list_to_doc(doc, list_items)
                list_items = []
                in_list = False
            
            if in_code_block:
                # Add code line with monospace font
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Courier New'
            elif line.strip():
                # Regular paragraph with basic formatting
                text = _process_inline_formatting(line)
                if text.strip():
                    _add_formatted_paragraph(doc, text)
        
        i += 1
    
    # Handle any remaining list items
    if in_list and list_items:
        _add_list_to_doc(doc, list_items)
    
    # Save to file if output path provided
    if output_path:
        doc.save(output_path)
        return output_path
    
    return doc


def _add_list_to_doc(doc, items):
    """Add a bulleted list to the document."""
    for item in items:
        p = doc.add_paragraph()
        p.style = 'List Bullet'
        _add_formatted_text_to_paragraph(p, item)


def _process_inline_formatting(text: str) -> str:
    """Process basic inline formatting like bold and italic."""
    return text


def _add_formatted_paragraph(doc, text: str):
    """Add a paragraph with basic formatting support."""
    p = doc.add_paragraph()
    _add_formatted_text_to_paragraph(p, text)


def _add_formatted_text_to_paragraph(paragraph, text: str):
    """Add formatted text to a paragraph, handling bold and italic."""
    # Simple implementation - can be enhanced for more complex formatting
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(part)